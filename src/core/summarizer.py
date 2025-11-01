"""
Core Summarization Service
Handles text processing, AI model integration, and export functionality
"""

import asyncio
import time
import os
import tempfile
from typing import Dict, Any, Literal
from pathlib import Path

# AI Model imports
from transformers import T5Tokenizer, T5ForConditionalGeneration
import torch

# Export libraries
from fpdf import FPDF  # fpdf2 package
from docx import Document
from docx.shared import Inches
import markdown

class SummarizerService:
    def __init__(self):
        self.tokenizer = None
        self.model = None
        self.model_loaded = False
        
    async def load_model(self):
        """Load the T5 model asynchronously - optimized for free tier"""
        if self.model_loaded:
            return
            
        try:
            print("[INFO] Loading T5 model (this may take 30-60 seconds on free tier)...")
            print("[INFO] Tip: First request after spin-down will be slower due to model loading")
            
            # Run model loading in thread pool to avoid blocking
            loop = asyncio.get_event_loop()
            self.tokenizer, self.model = await loop.run_in_executor(
                None, self._load_model_sync
            )
            self.model_loaded = True
            print("[OK] T5 model loaded successfully!")
        except Exception as e:
            print(f"[ERROR] Error loading model: {e}")
            print("[INFO] Falling back to extractive summarization only")
            self.model_loaded = False
    
    def _load_model_sync(self):
        """Synchronous model loading - optimized for free tier"""
        model_name = "t5-small"
        
        # Optimize for free tier memory constraints
        tokenizer = T5Tokenizer.from_pretrained(
            model_name,
            cache_dir="./model_cache",
            local_files_only=False
        )
        
        model = T5ForConditionalGeneration.from_pretrained(
            model_name,
            cache_dir="./model_cache",
            local_files_only=False,
            torch_dtype=torch.float32,  # Use float32 instead of float16 for compatibility
            low_cpu_mem_usage=True
        )
        
        # Set model to evaluation mode and optimize memory
        model.eval()
        model.config.use_cache = False
        
        return tokenizer, model
    
    def _tokenize_length(self, text: str) -> int:
        """Count the number of tokens in a text for T5 tokenizer limit check."""
        if not self.tokenizer:
            # Fallback guess: 1 word = 1 token
            return len(text.split())
        return len(self.tokenizer.encode(text, truncation=False))

    def _split_into_chunks(self, text: str, max_tokens: int = 512) -> list:
        """Splits text into chunks each <= max_tokens for T5 input limit."""
        if not self.tokenizer:
            # Fallback: naive split by word count
            words = text.split()
            return [" ".join(words[i:i+max_tokens]) for i in range(0, len(words), max_tokens)]
        tokens = self.tokenizer.encode(text, truncation=False)
        chunks = [tokens[i:i+max_tokens] for i in range(0, len(tokens), max_tokens)]
        return [self.tokenizer.decode(chunk, skip_special_tokens=True) for chunk in chunks]

    async def summarize_text(self, text: str, language: Literal["hindi", "english"] = "hindi", summary_length: Literal["short", "medium", "long", "auto"] = "auto") -> Dict[str, Any]:
        """Summarize text using AI model; universal chunking/truncation for huge inputs."""
        try:
            start_time = time.time()
            
            # Validate input
            if not text or not text.strip():
                raise ValueError("Text cannot be empty")
            
            # Handle auto summary length
            if summary_length == "auto":
                word_count = len(text.split())
                if word_count < 100:
                    summary_length = "short"
                elif word_count < 500:
                    summary_length = "medium"
                else:
                    summary_length = "long"
            
            # For Hindi, prefer extractive summarization (T5-small not reliable for Hindi)
            # For small texts, use extractive directly (works better for short/medium inputs)
            word_count = len(text.split())
            if word_count < 1000 or language == "hindi":  # For smaller texts OR Hindi, extractive is more reliable
                return await self._extractive_summarize(text, language, summary_length)
            
            # For large English texts only, use chunking + neural model
            await self.load_model()
            max_tokens = 512
            
            # Determine max summary length per chunk based on summary_length
            if summary_length == "short":
                chunk_max_summary = 50
            elif summary_length == "medium":
                chunk_max_summary = 100
            else:  # long
                chunk_max_summary = 200
            
            chunks = self._split_into_chunks(text, max_tokens=max_tokens)
            summary_chunks = []
            
            for ch in chunks:
                # Skip if chunk too small
                if self._tokenize_length(ch) < 10:
                    continue
                
                # Get summary for chunk
                if self.model_loaded:
                    # Use neural model (only for English)
                    prompt = "summarize: " + ch
                    input_ids = self.tokenizer.encode(prompt, return_tensors="pt", truncation=True, max_length=max_tokens)
                    summary_ids = self.model.generate(input_ids, max_length=chunk_max_summary, num_beams=4, early_stopping=True)
                    summary = self.tokenizer.decode(summary_ids[0], skip_special_tokens=True)
                    # Remove the prompt prefix if present
                    if summary.startswith("summarize: "):
                        summary = summary.replace("summarize: ", "", 1)
                    # Validate model output - if it looks like garbage (too many special chars/numbers), use extractive
                    special_char_ratio = sum(1 for c in summary if c in '[(){}:,-]') / max(len(summary), 1)
                    if special_char_ratio > 0.3 or len(summary.split()) < 3:  # Too many special chars or too short
                        summary_result = await self._extractive_summarize(ch, language, summary_length)
                        summary = summary_result['summary']
                else:
                    # Fallback to extractive - use summary_length parameter!
                    summary_result = await self._extractive_summarize(ch, language, summary_length)
                    summary = summary_result['summary']
                
                if summary and summary.strip():
                    summary_chunks.append(summary.strip())
            
            # Combine summaries - if too many chunks, limit final length
            if len(summary_chunks) > 1:
                combined_summary = "\n".join(summary_chunks)
                # If combined is too long, take first N chunks based on summary_length
                combined_words = len(combined_summary.split())
                if summary_length == "short" and combined_words > 100:
                    # Take first few sentences
                    sentences = combined_summary.split('.')
                    summary = '. '.join(sentences[:5]) + '.'
                elif summary_length == "medium" and combined_words > 200:
                    sentences = combined_summary.split('.')
                    summary = '. '.join(sentences[:10]) + '.'
                elif summary_length == "long" and combined_words > 300:
                    # Truncate long summaries to max 300 words at sentence boundaries
                    sentences = combined_summary.split('.')
                    summary_words = 0
                    truncated_sentences = []
                    for sentence in sentences:
                        sentence_words = len(sentence.split())
                        if summary_words + sentence_words <= 300:
                            truncated_sentences.append(sentence)
                            summary_words += sentence_words
                        else:
                            break
                    summary = '. '.join(truncated_sentences) + '.' if truncated_sentences else combined_summary[:500] + '...'
                else:
                    summary = combined_summary
            elif summary_chunks:
                summary = summary_chunks[0]
            else:
                # Fallback to extractive if all chunks failed
                return await self._extractive_summarize(text, language, summary_length)
            
            processing_time = time.time() - start_time
            summary_words = len(summary.split())
            
            return {
                "summary": summary,
                "original_length": word_count,
                "summary_length": summary_words,
                "compression_ratio": round(summary_words / word_count, 2) if word_count else 0.0,
                "processing_time": round(processing_time, 2)
            }
        except Exception as e:
            print(f"Error in summarize_text (universal chunking): {e}")
            # Fallback to extractive on any error
            try:
                return await self._extractive_summarize(text, language, summary_length if summary_length != "auto" else "medium")
            except:
                raise Exception(f"Failed to summarize text: {str(e)}")
    
    async def _extractive_summarize(
        self, 
        text: str, 
        language: Literal["hindi", "english"],
        summary_length: Literal["short", "medium", "long"]
    ) -> Dict[str, Any]:
        """Improved extractive summarization"""
        start_time = time.time()
        
        # Better sentence splitting for both English and Hindi
        import re
        # Handle both English and Hindi sentence endings
        sentences = re.split(r'[.!?।]+', text)
        sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 10]
        
        word_count = len(text.split())
        
        # Better length mapping based on word count
        if summary_length == "short":
            target_words = min(50, max(20, word_count // 8))
        elif summary_length == "medium":
            target_words = min(100, max(40, word_count // 5))
        else:  # long
            target_words = min(200, max(80, word_count // 3))
        
        # Select sentences to reach target word count
        summary_sentences = []
        current_words = 0
        
        for sentence in sentences:
            sentence_words = len(sentence.split())
            if current_words + sentence_words <= target_words:
                summary_sentences.append(sentence)
                current_words += sentence_words
            else:
                # Add partial sentence if we're close to target
                if current_words < target_words * 0.7:  # If we're less than 70% of target
                    remaining_words = target_words - current_words
                    words = sentence.split()
                    if len(words) > remaining_words:
                        partial_sentence = ' '.join(words[:remaining_words])
                        summary_sentences.append(partial_sentence + "...")
                    else:
                        summary_sentences.append(sentence)
                break
        
        # If no sentences selected, take first sentence
        if not summary_sentences:
            summary_sentences = [sentences[0]] if sentences else [text[:100] + "..."]
        
        summary = '. '.join(summary_sentences).strip()
        
        # Ensure summary ends with proper punctuation
        if summary and not summary.endswith(('.', '!', '?')):
            summary += '.'
        
        processing_time = time.time() - start_time
        
        return {
            "summary": summary,
            "original_length": word_count,
            "summary_length": len(summary.split()),
            "compression_ratio": round(len(summary.split()) / word_count, 2),
            "processing_time": round(processing_time, 2)
        }
    
    async def summarize_url(
        self, 
        url: str, 
        language: Literal["hindi", "english"] = "hindi",
        summary_length: Literal["short", "medium", "long", "auto"] = "auto"
    ) -> Dict[str, Any]:
        """Extract and summarize content from URL"""
        try:
            # Validate URL
            if not url or not url.strip():
                raise ValueError("URL cannot be empty")
            
            if not url.startswith(('http://', 'https://')):
                url = 'https://' + url
            
            from newspaper import Article
            
            # Extract article content
            article = Article(url, language="hi" if language == "hindi" else "en")
            article.download()
            article.parse()
            
            text = article.text.strip()
            title = article.title or "Untitled Article"
            
            if not text:
                raise ValueError("No content extracted from URL. The website might not be accessible or doesn't contain readable text.")
            
            # Summarize the extracted text
            result = await self.summarize_text(text, language, summary_length)
            result["title"] = title
            result["url"] = url
            
            return result
            
        except Exception as e:
            print(f"Error in summarize_url: {e}")
            raise Exception(f"Failed to process URL: {str(e)}")
    
    async def export_pdf(
        self, 
        summary: str, 
        title: str = "Summary",
        language: Literal["hindi", "english"] = "hindi"
    ) -> str:
        """Export summary as PDF with proper font support for Hindi"""
        # Fix Windows console encoding at function start
        import sys
        if sys.platform == 'win32' and hasattr(sys.stdout, 'buffer'):
            try:
                import io
                if not isinstance(sys.stdout, io.TextIOWrapper) or sys.stdout.encoding != 'utf-8':
                    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
            except:
                pass
        
        try:
            pdf = FPDF()
            pdf.add_page()  # REQUIRED: Must call add_page() before any operations
            
            # Handle Hindi vs English differently
            if language == "hindi":
                # HINDI EXPORT - Use Unicode font
                font_loaded = False
                try:
                    # Get font path - try multiple possible locations
                    font_path = Path(__file__).parent.parent.parent / "fonts" / "NotoSansDevanagari-Regular.ttf"
                    if not font_path.exists():
                        # Try alternative path
                        font_path = Path.cwd() / "fonts" / "NotoSansDevanagari-Regular.ttf"
                    
                    if font_path.exists():
                        # Register both regular and bold variants for Hindi font
                        pdf.add_font('NotoSansDevanagari', '', str(font_path))
                        pdf.add_font('NotoSansDevanagari', 'B', str(font_path))
                        pdf.set_font("NotoSansDevanagari", "B", 16)
                        font_loaded = True
                        print(f"[INFO] Hindi font loaded successfully from: {font_path}")
                        print(f"[INFO] Font file size: {font_path.stat().st_size} bytes")
                        print(f"[INFO] Summary text length: {len(summary)} chars")
                        # Avoid printing Hindi text directly to prevent encoding errors on Windows
                        try:
                            print(f"[INFO] First 50 chars (bytes): {summary[:50].encode('utf-8', errors='ignore')}")
                        except:
                            print(f"[INFO] Summary preview: {len(summary[:50])} chars")
                    else:
                        print(f"[WARN] Hindi font not found at: {font_path}")
                        raise FileNotFoundError(f"Font not found: {font_path}")
                except Exception as font_error:
                    # Safe error printing for Windows console
                    err_str = str(font_error).encode('ascii', errors='replace').decode('ascii')
                    print(f"[ERROR] Could not load Hindi font: {err_str}")
                    try:
                        import traceback
                        traceback.print_exc()
                    except:
                        pass
                    # Fallback to Helvetica
                    pdf.set_font("Helvetica", "B", 16)
                    font_loaded = False
                
                # Add title with Hindi font
                # Note: Font is already set above if font_loaded=True
                try:
                    if not font_loaded:
                        pdf.set_font("Helvetica", "B", 16)
                    pdf.cell(0, 10, title, 0, 1, "C")
                except Exception as title_err:
                    # If title addition fails, use fallback
                    try:
                        err_msg = str(title_err).encode('ascii', errors='replace').decode('ascii')
                        print(f"[WARN] Error adding Hindi title: {err_msg}")
                    except:
                        print("[WARN] Error adding Hindi title")
                    # Ensure page and font are set for fallback
                    try:
                        if not hasattr(pdf, 'page') or pdf.page == 0:
                            pdf.add_page()
                    except:
                        pdf.add_page()
                    pdf.set_font("Helvetica", "B", 16)
                    pdf.cell(0, 10, "Summary", 0, 1, "C")
                
                pdf.ln(10)
                
                # Set font for summary content
                if font_loaded:
                    try:
                        pdf.set_font("NotoSansDevanagari", "", 12)
                    except Exception as e:
                        print(f"[WARN] Failed to set Hindi font for content: {e}")
                        pdf.set_font("Helvetica", "", 12)
                        font_loaded = False
                else:
                    pdf.set_font("Helvetica", "", 12)
                
                # Add Hindi summary - fpdf2 should handle Unicode with the font
                try:
                    if font_loaded:
                        # For Hindi, split into smaller chunks and process carefully
                        paragraphs = summary.split('\n\n') if '\n\n' in summary else [summary]
                        for para in paragraphs:
                            if para.strip():
                                lines = para.split('\n')
                                for line in lines:
                                    if line.strip():
                                        try:
                                            pdf.multi_cell(0, 8, line.strip())
                                            pdf.ln(2)
                                        except Exception as line_err:
                                            # Avoid printing Hindi text in error (Windows encoding)
                                            err_msg = str(line_err).encode('ascii', errors='replace').decode('ascii')
                                            print(f"[WARN] Error adding Hindi line: {err_msg}")
                                            # Try with encoded version as last resort
                                            try:
                                                line_safe = line.strip().encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')
                                                pdf.multi_cell(0, 8, line_safe)
                                                pdf.ln(2)
                                            except:
                                                continue
                                    else:
                                        pdf.ln(4)
                                pdf.ln(4)
                    else:
                        # Fallback without font - will show question marks
                        pdf.multi_cell(0, 8, summary)
                except Exception as text_error:
                    print(f"[ERROR] Error adding Hindi summary text: {text_error}")
                    import traceback
                    traceback.print_exc()
                    # Try one more time with basic encoding
                    try:
                        summary_utf8 = summary.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')
                        pdf.multi_cell(0, 8, summary_utf8)
                    except:
                        pdf.multi_cell(0, 8, "Summary could not be encoded properly")
            else:
                # ENGLISH EXPORT - Use standard Helvetica
                # IMPORTANT: Helvetica doesn't support Unicode - sanitize text first
                import unicodedata
                
                # Sanitize title - replace Unicode chars with ASCII equivalents
                title_safe = title.encode('ascii', errors='replace').decode('ascii')
                # Also replace common Unicode punctuation with ASCII
                title_safe = title_safe.replace('\u2018', "'").replace('\u2019', "'")  # smart quotes
                title_safe = title_safe.replace('\u201C', '"').replace('\u201D', '"')  # smart double quotes
                title_safe = title_safe.replace('\u2013', '-').replace('\u2014', '--')  # en/em dashes
                title_safe = title_safe.replace('\u2026', '...')  # ellipsis
                
                try:
                    pdf.set_font("Helvetica", "B", 16)
                    pdf.cell(0, 10, title_safe, 0, 1, "C")
                except Exception as title_err:
                    # If still fails, use generic title
                    pdf.cell(0, 10, "Summary", 0, 1, "C")
                
                pdf.ln(10)
                pdf.set_font("Helvetica", "", 12)
                
                # Sanitize summary - replace Unicode chars with ASCII equivalents
                summary_safe = summary.encode('ascii', errors='replace').decode('ascii')
                # Replace common Unicode punctuation
                summary_safe = summary_safe.replace('\u2018', "'").replace('\u2019', "'")  # smart quotes
                summary_safe = summary_safe.replace('\u201C', '"').replace('\u201D', '"')  # smart double quotes
                summary_safe = summary_safe.replace('\u2013', '-').replace('\u2014', '--')  # en/em dashes
                summary_safe = summary_safe.replace('\u2026', '...')  # ellipsis
                
                # Add English summary
                try:
                    pdf.multi_cell(0, 8, summary_safe)
                except Exception as summary_err:
                    # If still fails, try even more aggressive sanitization
                    try:
                        err_str = str(summary_err).encode('ascii', errors='replace').decode('ascii')
                        print(f"[WARN] English PDF encoding error: {err_str}")
                        # Remove all non-ASCII characters
                        summary_stripped = ''.join(char for char in summary if ord(char) < 128)
                        pdf.multi_cell(0, 8, summary_stripped)
                    except Exception as fallback_err:
                        # Last resort
                        pdf.multi_cell(0, 8, "Summary contains characters that cannot be encoded in PDF format.")
            
            # Add footer (same for both languages)
            pdf.ln(20)
            pdf.set_font("Helvetica", "I", 8)
            footer_text = f"Generated by MultiLanguage AI Text Summarizer - {time.strftime('%Y-%m-%d %H:%M')}"
            pdf.cell(0, 10, footer_text, 0, 1, "C")
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            pdf.output(temp_file.name)
            
            return temp_file.name
            
        except Exception as e:
            
            # Catch ALL exceptions and convert to safe ASCII BEFORE any string operations
            # This prevents Windows encoding errors when exception messages contain Unicode
            try:
                # Get exception type name safely
                exc_type = type(e).__name__
                # Get exception message safely - convert to bytes first, then ASCII
                try:
                    exc_msg_bytes = repr(e).encode('utf-8', errors='replace')
                    exc_msg_safe = exc_msg_bytes.decode('ascii', errors='replace')
                except:
                    exc_msg_safe = "Unknown error (encoding issue)"
                
                print(f"[ERROR] PDF export error ({exc_type}): {exc_msg_safe[:200]}")
            except Exception as safe_err:
                # If even safe handling fails, use minimal message
                print("[ERROR] PDF export error: Could not encode error message")
            
            # Don't try to print traceback as it might contain Unicode
            # Raise a safe exception
            raise Exception("Failed to create PDF: Encoding or font issue with Hindi text")
    
    async def export_word(
        self, 
        summary: str, 
        title: str = "Summary",
        language: Literal["hindi", "english"] = "hindi"
    ) -> str:
        """Export summary as Word document with proper Hindi support"""
        try:
            doc = Document()
            
            # Handle Hindi vs English differently
            if language == "hindi":
                # HINDI EXPORT - Set Hindi-compatible font
                title_para = doc.add_heading(title, 0)
                
                # Try to set Hindi font - use fonts that commonly support Devanagari
                hindi_fonts = ['Mangal', 'Nirmala UI', 'Arial Unicode MS', 'Noto Sans Devanagari']
                font_set = False
                
                try:
                    from docx.shared import Pt
                    from docx.oxml.ns import qn
                    
                    # Set font for title
                    if title_para.runs:
                        for run in title_para.runs:
                            for font_name in hindi_fonts:
                                try:
                                    run.font.name = font_name
                                    run.font.size = Pt(16)
                                    font_set = True
                                    break
                                except:
                                    continue
                    else:
                        # No runs, create one
                        run = title_para.add_run(title)
                        for font_name in hindi_fonts:
                            try:
                                run.font.name = font_name
                                run.font.size = Pt(16)
                                font_set = True
                                break
                            except:
                                continue
                    
                    print(f"[INFO] Hindi Word export: Font {'set' if font_set else 'using default'}")
                except Exception as font_error:
                    print(f"[WARN] Could not set Hindi font for Word title: {font_error}")
                
                # Add summary paragraph
                summary_para = doc.add_paragraph(summary)
                
                # Set font for summary content
                try:
                    from docx.shared import Pt
                    for run in summary_para.runs:
                        if not font_set:
                            for font_name in hindi_fonts:
                                try:
                                    run.font.name = font_name
                                    run.font.size = Pt(12)
                                    font_set = True
                                    break
                                except:
                                    continue
                        else:
                            run.font.size = Pt(12)
                except Exception as font_error:
                    print(f"[WARN] Could not set Hindi font for Word content: {font_error}")
                
            else:
                # ENGLISH EXPORT - Use default fonts (unchanged from original)
                title_para = doc.add_heading(title, 0)
                summary_para = doc.add_paragraph(summary)
            
            # Add metadata (same for both)
            doc.add_paragraph(f"\n\nGenerated by MultiLanguage AI Text Summarizer")
            doc.add_paragraph(f"Date: {time.strftime('%Y-%m-%d %H:%M')}")
            
            # Save to temporary file - python-docx handles UTF-8 automatically
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            doc.save(temp_file.name)
            
            return temp_file.name
            
        except Exception as e:
            print(f"[ERROR] Word export error: {e}")
            import traceback
            traceback.print_exc()
            raise Exception(f"Failed to create Word document: {str(e)}")
    
    async def export_markdown(
        self, 
        summary: str, 
        title: str = "Summary",
        language: Literal["hindi", "english"] = "hindi"
    ) -> str:
        """Export summary as Markdown"""
        try:
            markdown_content = f"""# {title}

{summary}

---

*Generated by MultiLanguage AI Text Summarizer*  
*Date: {time.strftime('%Y-%m-%d %H:%M')}*
"""
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".md", mode='w', encoding='utf-8')
            temp_file.write(markdown_content)
            temp_file.close()
            
            return temp_file.name
            
        except Exception as e:
            raise Exception(f"Failed to create Markdown file: {str(e)}")
